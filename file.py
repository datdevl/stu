from docx import Document
from docx.shared import RGBColor

def process_docx(input_path, output_path):
    # Mở tài liệu .docx
    doc = Document(input_path)
    
    # Tạo tài liệu mới để lưu kết quả
    new_doc = Document()

    # Duyệt qua tất cả các đoạn văn trong tài liệu
    for para in doc.paragraphs:
        new_paragraph = ""
        add_paragraph = False
        
        # Duyệt qua tất cả các run (đoạn văn bản) trong mỗi paragraph
        for run in para.runs:
            # Kiểm tra màu chữ có phải là đỏ không (RGBColor(255, 0, 0) là đỏ)
            if run.font.color and run.font.color.rgb == RGBColor(255, 0, 0):
                # Nếu là đáp án đúng (màu đỏ), thêm vào đoạn văn mới
                new_paragraph += run.text
                add_paragraph = True
        
        # Nếu đoạn văn chứa đáp án đúng, thêm vào tài liệu mới
        if add_paragraph:
            new_doc.add_paragraph(new_paragraph)

    # Lưu tài liệu mới
    new_doc.save(output_path)
    print(f"Tài liệu đã được lưu tại: {output_path}")
# Đường dẫn đến file gốc và file đầu ra
input_file = "/home/levandat/Desktop/asinhsinh.docx"  # Đảm bảo đường dẫn đúng
output_file = "/home/levandat/Desktop/output_file.docx"  # Đảm bảo lưu đúng vị trí

# Gọi hàm để xử lý và lưu lại kết quả
process_docx(input_file, output_file)
